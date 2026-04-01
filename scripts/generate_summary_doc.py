"""
Generate a Word document summarizing the polycrisis research conversation findings.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import datetime


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading("Polycrisis & Limits to Growth\nIndependent Research Plan", level=0)
    doc.add_paragraph(f"Compiled: {datetime.date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph("Author: Alex Coffman")
    doc.add_paragraph("Research Assistant: Claude (Anthropic)")
    doc.add_paragraph()

    # ── SECTION 1: BACKGROUND ──
    add_heading(doc, "1. Background & Research Goals")
    doc.add_paragraph(
        "This document captures the findings from our initial research session planning "
        "independent academic research on polycrisis and limits-to-growth topics. The goal "
        "is to produce peer-reviewed publications and build a public research profile outside "
        "of a formal degree program."
    )

    add_heading(doc, "What Alex Asked For", level=2)
    bullets = [
        "Map every instance of 'polycrisis' and 'limits to growth' in academic literature and show publication trends over time",
        "Identify the top voices in the space and who would respond to outreach from an independent researcher with a Substack",
        "Identify gaps in the existing literature that can be studied from a computer (no proprietary data, no physical experiments)",
        "Identify premier journals and publication venues for this research area",
        "Build a pipeline from Substack → preprint → peer-reviewed journal",
        "Start a live World3 dashboard project (ambitious, longer-term)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    add_heading(doc, "Alex's Constraints & Resources", level=2)
    constraints = [
        "No institutional affiliation (independent researcher)",
        "Free data access only (no Scopus/Web of Science subscription)",
        "Python skills via AI assistance (Claude + Codex)",
        "Existing Substack at acoffman.substack.com (~10 posts, casual voice, multi-topic)",
        "Timeline: ASAP for first publication, then build toward bigger projects",
    ]
    for c in constraints:
        doc.add_paragraph(c, style="List Bullet")

    # ── SECTION 2: PUBLICATION TRENDS ──
    add_heading(doc, "2. Publication Trend Data")
    doc.add_paragraph(
        "Data pulled live from the OpenAlex API (api.openalex.org), the largest free academic "
        "database with 250M+ works. No API key required."
    )

    add_heading(doc, "Polycrisis Papers by Year (OpenAlex)", level=2)
    add_table(doc,
        ["Year", "Papers", "YoY Growth"],
        [
            ["2016", "2", "--"],
            ["2017", "1", "-50%"],
            ["2018", "6", "+500%"],
            ["2019", "10", "+67%"],
            ["2020", "13", "+30%"],
            ["2021", "10", "-23%"],
            ["2022", "38", "+280%"],
            ["2023", "136", "+258%"],
            ["2024", "358/361", "+163%"],
            ["2025", "672/677", "+88%"],
            ["2026 (partial)", "227", "--"],
        ]
    )
    doc.add_paragraph(
        "Key finding: Near-zero before 2022, then exponential. The inflection was created by "
        "Adam Tooze's Financial Times column, the Cascade Institute's definition paper, and "
        "the World Economic Forum's 2023 Global Risks Report. At current trajectory, expect "
        "1,000+ papers in 2026. The field is still in growth phase — not saturated."
    )

    add_heading(doc, "'Limits to Growth' (exact phrase) by Year", level=2)
    add_table(doc,
        ["Year", "Papers", "YoY Growth"],
        [
            ["2020", "599", "--"],
            ["2021", "625", "+4%"],
            ["2022", "683", "+9%"],
            ["2023", "709", "+4%"],
            ["2024", "723", "+2%"],
            ["2025", "927", "+28%"],
        ]
    )
    doc.add_paragraph(
        "Steady ~600-700/year with a 28% jump in 2025, likely driven by Club of Rome 50th "
        "anniversary and convergence with polycrisis discourse."
    )
    doc.add_paragraph(
        "DATA QUALITY NOTE: The LtG query returns many biology papers ('iron limits to growth "
        "of phytoplankton,' cancer cell growth limits). The top cited results include unrelated "
        "Nature papers. A refined query filtering by OpenAlex concept tags or co-author names "
        "is needed before formal bibliometric analysis."
    )

    add_heading(doc, "Key Inflection Points", level=2)
    inflections = [
        "2008: Turner's validation reignited LtG after decades of dismissal",
        "2021: Herrington's KPMG paper went viral — first mainstream LtG coverage since the 1970s",
        "2022: Tooze + WEF + Cascade Institute created the 'polycrisis moment'",
        "2023-2025: Exponential growth; systematic reviews arriving (field maturing)",
        "2024-2025: LtG and polycrisis are CONVERGING — researchers connecting resource limits to cascading crises",
    ]
    for i in inflections:
        doc.add_paragraph(i, style="List Bullet")

    doc.add_paragraph(
        "Critical observation: The 2025 Annual Review systematic review scanned 2,299 publications "
        "but only analyzed 59 in depth. ~2,300 papers touch polycrisis but only ~60 are rigorous, "
        "focused contributions. There is a quality gap — and an opportunity."
    )

    # ── SECTION 3: JOURNALS ──
    add_heading(doc, "3. Premier Journals & Publication Venues")

    add_heading(doc, "Tier 1 — Primary Targets", level=2)
    add_table(doc,
        ["Journal", "Publisher", "Impact Factor", "Why"],
        [
            ["Global Sustainability", "Cambridge Univ Press", "~6.0",
             "THE polycrisis journal. Open-access. Active 'Polycrisis in the Anthropocene' collection. Michael Lawrence is Section Editor."],
            ["Journal of Industrial Ecology", "Wiley (Yale)", "~5.9",
             "Where Herrington (2021) and Nebel (2024) published LtG validation/recalibration work."],
            ["Sustainability (MDPI)", "MDPI", "~3.9",
             "Faster review, more accessible. Good for first publication. Open-access."],
        ]
    )

    add_heading(doc, "Tier 2 — Higher Impact", level=2)
    add_table(doc,
        ["Journal", "Publisher", "Impact Factor", "Why"],
        [
            ["Annual Review of Env & Resources", "Annual Reviews", "~16.0", "2025 systematic review published here"],
            ["Nature Sustainability", "Nature", "~27.0", "Top-tier, very competitive"],
            ["Global Environmental Change", "Elsevier", "~11.0", "Where Turner (2008) published"],
            ["PNAS Nexus", "NAS", "New", "Open-access, accepts novel methods"],
        ]
    )

    add_heading(doc, "Active Calls for Papers", level=2)
    cfps = [
        "Global Sustainability — 'Polycrisis in the Anthropocene' (Cambridge). Michael Lawrence is Section Editor. No university affiliation required. HIGHEST VALUE TARGET.",
        "Journal of Economic Geography — 'Restructuring and Resilience of Global Production Networks in the Age of Polycrisis' (Oxford).",
        "Springer Nature — Open Polycrisis Collection.",
        "Int. Journal of Disaster Risk Science — 2025 special issue on Polycrisis and Systemic Risks.",
    ]
    for c in cfps:
        doc.add_paragraph(c, style="List Bullet")

    add_heading(doc, "Key Conferences", level=2)
    add_table(doc,
        ["Conference", "When", "Where"],
        [
            ["Global Tipping Points", "Jun 30 - Jul 3, 2025", "Exeter, UK (hybrid)"],
            ["ISEE + Degrowth", "Jun 24-27, 2025", "Oslo, Norway"],
            ["System Dynamics Society (ISDC)", "Aug 3-7, 2025; Jul 20-24, 2026", "Boston; TU Delft"],
            ["USSEE Biennial", "Jun 18-21, 2026", "Oberlin College, Ohio"],
            ["Club of Rome", "2025 Suzhou; Jan 2026 Reclaim Economy", "Various"],
        ]
    )

    # ── SECTION 4: RESEARCHERS ──
    add_heading(doc, "4. Researcher Accessibility Ranking")
    doc.add_paragraph(
        "Ranked by likelihood of engaging with an independent researcher who has a "
        "Substack and a working paper."
    )

    add_heading(doc, "Tier 1 — Most Likely to Engage", level=2)
    add_table(doc,
        ["Researcher", "Institution", "Platform", "Why Accessible"],
        [
            ["Gaya Herrington", "Independent (Club of Rome, Harvard)", "LinkedIn, Bluesky, website",
             "Contact form EXPLICITLY lists 'research collaboration.' Independent-minded, left KPMG. Most-cited recent LtG validator."],
            ["Michael Lawrence", "Cascade Institute", "LinkedIn, Polycrisis.org",
             "Section Editor of Global Sustainability polycrisis collection. His job is to build the community. Low profile = less inbox noise."],
            ["Scott Janzwood", "Cascade Institute (Research Director)", "LinkedIn",
             "Former Future of Humanity Institute. Low profile = more bandwidth."],
        ]
    )

    add_heading(doc, "Tier 2 — Possible with Good Framing", level=2)
    add_table(doc,
        ["Researcher", "Institution", "Notes"],
        [
            ["Florian Jehn", "Independent", "Runs Existential Crunch Substack (~1K subs). Fellow indie voice. DM via Substack."],
            ["Jonathan Donges", "PIK + Stockholm Resilience Centre", "Open-source COPAN code on GitHub. Needs technically relevant pitch."],
            ["Nico Wunderling", "Goethe Univ Frankfurt", "Early-career professor, still building network. Very technical."],
        ]
    )

    add_heading(doc, "Tier 3 — Long Game", level=2)
    add_table(doc,
        ["Researcher", "Institution", "Notes"],
        [
            ["Adam Tooze", "Columbia", "221K Twitter, 181K Substack. Engage in Chartbook comments, don't cold email."],
            ["Thomas Homer-Dixon", "Cascade Institute", "Reach via Lawrence/Janzwood first. Warm intro target."],
        ]
    )

    add_heading(doc, "Recommended Outreach Sequence", level=2)
    steps = [
        "Build Substack with 2-3 substantive polycrisis posts",
        "Cold-email Gaya Herrington via gayaherrington.com/contact with a specific pitch",
        "Submit working paper to Global Sustainability (where Lawrence is editor)",
        "Engage in Chartbook (Tooze) and Existential Crunch (Jehn) comments",
        "Connect with Janzwood and Lawrence on LinkedIn after a published preprint",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    # ── SECTION 5: SUBSTACKS ──
    add_heading(doc, "5. Existing Substacks & Newsletters in the Space")
    add_table(doc,
        ["Newsletter", "Author", "Subscribers", "Focus"],
        [
            ["Chartbook", "Adam Tooze", "181K+", "Economics, geopolitics, polycrisis"],
            ["The Great Simplification", "Nate Hagens", "13K+", "Energy, ecology, overshoot"],
            ["The Honest Sorcerer", "Anonymous", "8.9K+", "Systems thinking, collapse"],
            ["Existential Crunch", "Florian Jehn", "1K+", "Existential risk, food security, polycrisis"],
            ["Degrowth is the Answer", "Matt Orsagh", "Thousands", "Degrowth, ecological limits"],
            ["The Back Loop", "Michael James", "New", "Post-doom, resilience"],
        ]
    )

    # ── SECTION 6: RESEARCH GAPS ──
    add_heading(doc, "6. Research Gaps Identified")

    add_heading(doc, "From the 2025 Annual Review Systematic Review", level=2)
    gaps1 = [
        "Crisis INTERACTIONS are underexplored — most papers study individual crises, not causal entanglement",
        "Lack of QUANTITATIVE empirical work — field is heavily theoretical/conceptual",
        "Systemic DRIVERS under-analyzed — inequality, extractive economies acknowledged but not formally modeled",
        "INTERVENTION evidence is thin — almost no studies evaluate which interventions reduce cascading",
        "Causal PATHWAYS need formal specification",
    ]
    for g in gaps1:
        doc.add_paragraph(g, style="List Bullet")

    add_heading(doc, "From the Cascade Institute 2024 Research Roadmap", level=2)
    gaps2 = [
        "Mechanisms of crisis TRANSMISSION between systems — top empirical priority",
        "Historical/comparative polycrisis analysis — almost absent",
        "Methods from complexity science, financial risk, network science are underused",
        "No systematic data infrastructure for polycrisis research",
    ]
    for g in gaps2:
        doc.add_paragraph(g, style="List Bullet")

    add_heading(doc, "Methods Available But Not Yet Applied to Polycrisis", level=2)
    methods = [
        "Multiplex network analysis (standard in financial contagion, never applied to polycrisis)",
        "Granger causality networks (only 1 study — Somalia — in climate-food-conflict nexus)",
        "Critical slowing down / early warning signals (mature in ecology, never tested on SDG indicators)",
        "Deep learning for tipping point detection (Bury et al. 2021, PNAS — never applied to socioeconomic data)",
        "Bibliometric mapping with VOSviewer/CiteSpace (no analysis of polycrisis field exists)",
    ]
    for m in methods:
        doc.add_paragraph(m, style="List Bullet")

    # ── SECTION 7: PROPOSED STUDIES ──
    add_heading(doc, "7. Proposed Research Projects")
    doc.add_paragraph("All projects use only publicly available data and can be done from a laptop.")

    add_heading(doc, "Quick Wins (2 months each)", level=2)

    doc.add_paragraph("Project 1: Bibliometric Mapping of Polycrisis Research", style="List Number")
    doc.add_paragraph("Gap: No bibliometric analysis exists. Method: OpenAlex API → VOSviewer co-citation/co-occurrence maps. "
                      "Target: Sustainability (MDPI). This is the recommended FIRST publication.")

    doc.add_paragraph("Project 2: SDG Regression Clustering", style="List Number")
    doc.add_paragraph("Gap: No systematic ID of countries regressing across multiple SDGs simultaneously. "
                      "Method: UN SDG data → k-means clustering → cross-reference with EM-DAT/ACLED/IMF. "
                      "Target: Global Sustainability or World Development.")

    add_heading(doc, "Medium Projects (3-4 months each)", level=2)

    doc.add_paragraph("Project 3: Granger Causality Networks of Crisis Transmission", style="List Number")
    doc.add_paragraph("Gap: Cascade Institute's #1 empirical priority. Method: Time series for 5-6 crisis domains "
                      "across 100+ countries → Granger causality → directed networks. Target: Global Sustainability.")

    doc.add_paragraph("Project 4: Early Warning Signals for Polycrisis", style="List Number")
    doc.add_paragraph("Gap: EWS methods never applied to polycrisis. Genuinely novel. Method: Rolling-window "
                      "autocorrelation/variance on SDG indicators → test predictive power. Target: PNAS Nexus.")

    doc.add_paragraph("Project 5: Climate-Food-Conflict Panel VAR", style="List Number")
    doc.add_paragraph("Gap: Only 1 Granger study exists (Somalia). Method: Panel VAR for 50-80 countries "
                      "with NOAA/FAO/ACLED data. Target: Global Environmental Change.")

    add_heading(doc, "Ambitious Projects (5-6 months each)", level=2)

    doc.add_paragraph("Project 6: Multiplex Network Model of Crisis Cascading", style="List Number")
    doc.add_paragraph("Gap: Never applied to polycrisis. Method: Countries as nodes, crisis domains as layers, "
                      "simulate cascading failures. Target: Nature Communications.")

    doc.add_paragraph("Project 7: Live World3 Dashboard with Real-Time Data", style="List Number")
    doc.add_paragraph("Gap: Nothing like this exists publicly. Method: pyworld3 + Nebel 2024 params + "
                      "World Bank/FAO/NOAA API feeds + Streamlit. Shows 'you are here' on LtG scenarios. "
                      "~1,500-2,800 LOC, 4-8 weeks. Target: Journal of Industrial Ecology.")

    add_heading(doc, "Recommended Sequence", level=2)
    doc.add_paragraph("1. Start with Project 1 (bibliometric mapping) — fast, guaranteed publishable, builds field knowledge")
    doc.add_paragraph("2. Then Project 2 (SDG regression clustering) — fast, policy-relevant, establishes empirical credibility")
    doc.add_paragraph("3. Then Project 3 or 4 as main contribution — fills the most-cited gaps")
    doc.add_paragraph("4. Project 7 (World3 dashboard) in parallel as a public-facing tool")

    # ── SECTION 8: DATASETS ──
    add_heading(doc, "8. Publicly Available Datasets")
    add_table(doc,
        ["Dataset", "URL", "Content"],
        [
            ["OpenAlex", "api.openalex.org", "250M+ academic works (free, no key)"],
            ["World Bank WDI", "data.worldbank.org", "1,400+ indicators, 200+ countries"],
            ["UN SDG Database", "unstats.un.org/sdgs/dataportal", "All SDG indicators"],
            ["EM-DAT", "emdat.be", "27,000+ disasters since 1900"],
            ["FAO Food Price Index", "fao.org", "Monthly commodity prices since 1990"],
            ["IMF WEO", "imf.org", "GDP, inflation, 190+ countries"],
            ["ACLED", "acleddata.com", "Political violence/protest events (free registration)"],
            ["NOAA", "ncei.noaa.gov", "Temperature, CO2, precipitation"],
            ["V-Dem", "v-dem.net", "Democracy indices, 200+ countries, 1789-present"],
            ["GDELT", "gdeltproject.org", "Media event data, daily since 1979"],
            ["World Inequality DB", "wid.world", "Income/wealth inequality, 100+ countries"],
        ]
    )

    # ── SECTION 9: WORLD3 DASHBOARD SCOPE ──
    add_heading(doc, "9. World3 Dashboard — Technical Scope")

    add_heading(doc, "Existing Code", level=2)
    add_table(doc,
        ["Repo", "Stars", "Notes"],
        [
            ["cvanwynsberghe/pyworld3", "339", "Gold standard. pip-installable, clean API, 1972 model"],
            ["TimSchell98/PyWorld3-03", "42", "Nebel 2024 recalibration. World3-03 (2004 update). CeCILL license."],
            ["worlddynamics/WorldDynamics.jl", "73", "Julia. Includes World3 + Earth4All"],
        ]
    )

    add_heading(doc, "Competition", level=2)
    doc.add_paragraph(
        "NONE. world3simulator.org is dead. En-ROADS is climate-only (not World3). "
        "Earth4All is Julia-only, not data-fed. No public live dashboard running actual "
        "World3 with real-time data exists anywhere. This project has strong novelty."
    )

    add_heading(doc, "Technical Estimate", level=2)
    add_table(doc,
        ["Component", "Lines of Code"],
        [
            ["World3 model integration", "200-400"],
            ["Data pipeline (5 APIs → normalized)", "500-800"],
            ["Scenario runner", "150-300"],
            ["Variable mapping (real data → World3)", "200-400"],
            ["Streamlit frontend", "400-700"],
            ["Caching + refresh logic", "100-200"],
            ["TOTAL", "~1,500-2,800"],
        ]
    )
    doc.add_paragraph(
        "Hardest part: Variable mapping. World3 variables like 'nonrenewable resources' or "
        "'industrial capital' are theoretical constructs with no direct real-world equivalent. "
        "Herrington (2021) spent most of her effort justifying proxy data series. This is the "
        "intellectual core of the project."
    )
    doc.add_paragraph("Estimated build time: 4-8 weeks at ~20hrs/week.")

    # ── SECTION 10: SUBSTACK STRATEGY ──
    add_heading(doc, "10. Substack Strategy")

    doc.add_paragraph(
        "Alex's existing Substack (acoffman.substack.com) covers AI risk, federal employment, "
        "personal finance, and systemic risk. The polycrisis topic fits naturally — his existing "
        "posts already demonstrate systems thinking (AI bubble analysis, DOGE governance erosion, "
        "perception-reality feedback loops)."
    )

    add_heading(doc, "Pipeline: Substack → Journal", level=2)
    pipeline = [
        "Substack post — workshop ideas, build audience, get informal feedback",
        "ArXiv/SSRN preprint — establish priority, get DOI, make citable",
        "Peer-reviewed journal — formal credibility, academic record",
        "Conference presentation — network with researchers, refine arguments",
    ]
    for i, p in enumerate(pipeline, 1):
        doc.add_paragraph(f"{i}. {p}")

    doc.add_paragraph(
        "IMPORTANT: Substack articles are NOT citable in academic contexts (no DOI, no peer review). "
        "Use Substack to develop ideas, then formalize via SSRN/ArXiv as intermediate step."
    )

    # ── SECTION 11: WHAT WE BUILT ──
    add_heading(doc, "11. What Was Built This Session")
    built = [
        "GitHub repo: github.com/Alzxcvb/polycrisis-research",
        "01_fetch_openalex.py — pulls all papers from OpenAlex API (1,500 polycrisis + 13,172 LtG papers fetched)",
        "02_bibliometric_analysis.py — trend charts, keyword co-occurrence, author/institution rankings, citation analysis",
        "03_vosviewer_export.py — formats data for VOSviewer visual network mapping",
        "Publication trend charts (figures/01_publication_trends.png)",
        "Journal distribution charts (figures/02_polycrisis_journals.png, 03_ltg_journals.png)",
        "Keyword co-occurrence CSVs ready for VOSviewer import",
        "Full research plan document: polycrisis-independent-research.md",
    ]
    for b in built:
        doc.add_paragraph(b, style="List Bullet")

    # ── SECTION 12: NEXT STEPS ──
    add_heading(doc, "12. Immediate Next Steps")
    next_steps = [
        "Download VOSviewer (vosviewer.com) and import polycrisis_keyword_cooccurrence.csv to see the cluster map — this becomes Figure 1 of the bibliometric paper",
        "Fix LtG query to exclude biology papers (add concept filters for Meadows/Club of Rome)",
        "Write first Substack post: 'What 1,500 papers tell us about the polycrisis' — share the trend data and cluster map",
        "Start World3 dashboard project: set up repo, begin variable mapping design based on Herrington (2021)",
        "Cold-email Gaya Herrington via her contact form once you have 1-2 substantial Substack posts",
        "Submit bibliometric working paper to SSRN/ArXiv for DOI, then target Sustainability (MDPI) for peer review",
    ]
    for i, s in enumerate(next_steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    # Save
    out_path = Path(__file__).parent.parent / "Polycrisis_Research_Summary.docx"
    doc.save(str(out_path))
    print(f"Document saved: {out_path}")


if __name__ == "__main__":
    main()
